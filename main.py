import asyncio
import os
import io
import json
import uuid
import random
from contextlib import suppress

# --- QO'SHIMCHA KUTUBXONALAR ---
from dotenv import load_dotenv
load_dotenv() 

import asyncpg
import google.generativeai as genai
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command, CommandObject, StateFilter
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove, InlineKeyboardMarkup, InlineKeyboardButton, BufferedInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

import PyPDF2
from docx import Document
from PIL import Image
from keep_alive import keep_alive

# --- SOZLAMALAR ---
BOT_TOKEN = os.getenv("BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

ADMIN_ID = 5031441892  # <--- DIQQAT: O'ZINGIZNING TELEGRAM ID RAQAMINGIZNI SHU YERGA YOZING!

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-2.5-flash')

# --- DATA ENGINE (Ma'lumotlar bazasi) ---
db_pool = None

async def init_db_pool():
    global db_pool
    db_pool = await asyncpg.create_pool(DATABASE_URL, min_size=5, max_size=20, statement_cache_size=0)
    async with db_pool.acquire() as conn:
        await conn.execute('''CREATE TABLE IF NOT EXISTS users (user_id TEXT PRIMARY KEY, name TEXT, score INTEGER DEFAULT 0, tests_taken INTEGER DEFAULT 0)''')
        await conn.execute('''CREATE TABLE IF NOT EXISTS quizzes (quiz_id TEXT PRIMARY KEY, source_type TEXT, savollar TEXT)''')
        
        columns_to_add = ["image_tests_made", "file_tests_made", "topic_tests_made"]
        for col in columns_to_add:
            try: await conn.execute(f"ALTER TABLE users ADD COLUMN {col} INTEGER DEFAULT 0")
            except Exception: pass
            
        try: await conn.execute("ALTER TABLE quizzes ADD COLUMN source_type TEXT")
        except Exception: pass
        try: await conn.execute("ALTER TABLE quizzes ADD COLUMN timer INTEGER DEFAULT 45")
        except Exception: pass
        
        try: await conn.execute("ALTER TABLE users ADD COLUMN phone_number TEXT")
        except Exception: pass

        try: await conn.execute("ALTER TABLE users ADD COLUMN joined_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
        except Exception: pass

async def add_user(user_id, name):
    async with db_pool.acquire() as conn:
        await conn.execute("INSERT INTO users (user_id, name) VALUES ($1, $2) ON CONFLICT (user_id) DO NOTHING", str(user_id), name)

POLL_DATA = {}   
SESSION_SCORES = {} 
USER_EVENTS = {}
ACTIVE_TESTS = {}

# --- YORDAMCHI FUNKSIYALAR ---
def clean_json_text(text):
    try:
        start = text.find('[')
        end = text.rfind(']')
        if start != -1 and end != -1: return text[start:end+1]
    except Exception: pass
    return text.strip()

def read_file_sync(file_data, filename):
    text = ""
    try:
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file_data)
            for page in pdf_reader.pages: text += page.extract_text() or ""
        elif filename.endswith('.docx'):
            doc = Document(file_data)
            for paragraph in doc.paragraphs: text += paragraph.text + "\n"
    except Exception as e: print(f"Fayl xatosi: {e}")
    return text

# --- FSM VA MENYULAR ---
class QuickQuizForm(StatesGroup):
    source_type = State()
    soni = State()
    til = State()
    vaqt = State()
    payload = State()
    filename = State()

class AdminState(StatesGroup):
    xabar_kutish = State()

class FeedbackState(StatesGroup):
    kutish = State()

bekor_tugma = [KeyboardButton(text="🔙 Bekor qilish")]

asosiy_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="📸 Rasmdan test"), KeyboardButton(text="📚 Matn/Mavzudan test")],
    [KeyboardButton(text="📊 Mening natijalarim"), KeyboardButton(text="🏆 Reyting")],
    [KeyboardButton(text="💬 Taklif va Xatolar")]
], resize_keyboard=True)

ixtiyoriy_raqam_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="📱 Raqamni ulashish", request_contact=True)],
    [KeyboardButton(text="➡️ Asosiy menyuga o'tish")]
], resize_keyboard=True, one_time_keyboard=True)

soni_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="15"), KeyboardButton(text="20")], 
    [KeyboardButton(text="25"), KeyboardButton(text="30")], 
    bekor_tugma
], resize_keyboard=True)

til_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="🇺🇿 O'zbek tili"), KeyboardButton(text="🇷🇺 Русский")],
    [KeyboardButton(text="🇬🇧 English")],
    bekor_tugma
], resize_keyboard=True)

vaqt_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="15 soniya"), KeyboardButton(text="30 soniya")], 
    [KeyboardButton(text="45 soniya"), KeyboardButton(text="60 soniya")], 
    bekor_tugma
], resize_keyboard=True)

admin_menyu = ReplyKeyboardMarkup(keyboard=[
    [KeyboardButton(text="📊 Umumiy Statistika"), KeyboardButton(text="👥 Foydalanuvchilar")],
    [KeyboardButton(text="📣 Xabar tarqatish"), KeyboardButton(text="🔙 Bosh menyu")]
], resize_keyboard=True)


# --- IXTIYORIY RAQAMNI QABUL QILISH ---
@dp.message(F.contact)
async def contact_handler(message: types.Message):
    telefon = message.contact.phone_number
    async with db_pool.acquire() as conn:
        await conn.execute("UPDATE users SET phone_number = $1 WHERE user_id = $2", telefon, str(message.from_user.id))
    await message.answer("✅ Rahmat! Telefon raqamingiz muvaffaqiyatli saqlandi.", reply_markup=asosiy_menyu)

@dp.message(F.text == "➡️ Asosiy menyuga o'tish")
async def skip_contact_handler(message: types.Message):
    await message.answer("Asosiy menyuga xush kelibsiz! Marhamat, xizmatlardan foydalaning.", reply_markup=asosiy_menyu)


# --- TAKLIF VA XATOLAR TIZIMI (FEEDBACK) ---
@dp.message(F.text == "💬 Taklif va Xatolar")
async def ask_feedback(message: types.Message, state: FSMContext):
    await message.answer("✍️ Bot bo'yicha qanday taklifingiz yoki topgan xatoligingiz bor? Marhamat, (matn, rasm, video yoki stiker) yuboring:", reply_markup=ReplyKeyboardMarkup(keyboard=[bekor_tugma], resize_keyboard=True))
    await state.set_state(FeedbackState.kutish)

# YANGILANGAN: Har qanday turdagi xabarni qabul qiladi
@dp.message(FeedbackState.kutish)
async def receive_feedback(message: types.Message, state: FSMContext):
    if message.text == "🔙 Bekor qilish":
        await state.clear()
        return await message.answer("Bekor qilindi.", reply_markup=asosiy_menyu)

    try:
        # 1. Adminga kimdan kelganini bildirish uchun pasport (ID) yuboramiz
        info_text = f"📬 **YANGI XABAR (Taklif/Xato)**\n👤 Kimdan: [{message.from_user.full_name}](tg://user?id={message.from_user.id})\n🆔 ID: `{message.from_user.id}`"
        await bot.send_message(ADMIN_ID, info_text, parse_mode="Markdown")
        
        # 2. Foydalanuvchi nima yuborgan bo'lsa (Stiker, Rasm, Matn), xuddi o'shani adminga nusxalaymiz
        await bot.copy_message(chat_id=ADMIN_ID, from_chat_id=message.chat.id, message_id=message.message_id)
        
        await message.answer("✅ Xabaringiz adminga muvaffaqiyatli yetkazildi! Fikringiz uchun rahmat.", reply_markup=asosiy_menyu)
    except Exception:
        await message.answer("⚠️ Adminga xabar yuborishda xatolik yuz berdi.", reply_markup=asosiy_menyu)
    await state.clear()

# YANGILANGAN: Admin har qanday formatda (hatto voice yoki rasm) javob yozishi mumkin
@dp.message(F.reply_to_message)
async def admin_reply_handler(message: types.Message):
    if message.from_user.id != int(ADMIN_ID): return
    
    # Biz faqat ID yozilgan matnga qilingan reply'ni qabul qilamiz
    original_text = message.reply_to_message.text
    if not original_text or "🆔 ID:" not in original_text: return
    
    try:
        target_id = original_text.split("🆔 ID: ")[1].split("\n")[0].strip('`')
        
        await bot.send_message(chat_id=target_id, text="👨‍💻 **Admindan javob:**", parse_mode="Markdown")
        # Admin nima yuborsa (stiker, voice, matn), foydalanuvchiga shuni nusxalab beramiz
        await bot.copy_message(chat_id=target_id, from_chat_id=message.chat.id, message_id=message.message_id)
        
        await message.answer("✅ Javobingiz foydalanuvchiga yuborildi!")
    except Exception as e:
        await message.answer(f"⚠️ Yuborishda xatolik: {e}")

# --- GURUHLARDA JONLI TEST (GROUP MODE) ---
@dp.message(Command("quiz"))
async def group_quiz_start(message: types.Message, command: CommandObject):
    if not command.args:
        return await message.answer("⚠️ Iltimos, test ID sini kiriting. Namuna: `/quiz 1234abcd`", parse_mode="Markdown")
        
    quiz_id = command.args
    async with db_pool.acquire() as conn:
        quiz_row = await conn.fetchrow("SELECT savollar, timer FROM quizzes WHERE quiz_id = $1", quiz_id)

    if not quiz_row:
        return await message.answer("⚠️ Bu test eskirgan yoki topilmadi.")

    taymer = quiz_row.get('timer') or 45
    savollar = json.loads(quiz_row['savollar'])

    await message.answer(f"🚀 **Guruh testi boshlanmoqda!**\nJami savollar: {len(savollar)} ta\nHar bir savolga: {taymer} soniya\nTo'xtatish uchun /stop bosing.\n\nTayyor turing!", parse_mode="Markdown")
    await asyncio.sleep(3) 

    ACTIVE_TESTS[message.chat.id] = True 
    
    for i, data in enumerate(savollar, 1):
        if not ACTIVE_TESTS.get(message.chat.id, True):
            break 

        q = f"[{i}/{len(savollar)}] {data['savol'][:200]}"
        opts = [str(opt)[:100] for opt in data['variantlar']][:4]
        correct = int(data.get('togri_index', 0))

        sent_poll = await bot.send_poll(chat_id=message.chat.id, question=q, options=opts, type='quiz', correct_option_id=correct, is_anonymous=False, open_period=taymer)
        POLL_DATA[sent_poll.poll.id] = {"correct": correct, "points": 2}
        await asyncio.sleep(taymer + 1)

    ACTIVE_TESTS[message.chat.id] = False
    await message.answer("🏁 **Guruh testi yakunlandi!**", parse_mode="Markdown")


# --- FAYL YUKLAB OLISH ---
@dp.callback_query(F.data.startswith("down_"))
async def download_doc(call: types.CallbackQuery):
    quiz_id = call.data.split("_")[1]
    async with db_pool.acquire() as conn:
        quiz_row = await conn.fetchrow("SELECT savollar FROM quizzes WHERE quiz_id = $1", quiz_id)
    if not quiz_row: return await call.answer("Test topilmadi!", show_alert=True)
        
    savollar = json.loads(quiz_row['savollar'])
    doc = Document()
    doc.add_heading(f"TestTuzar - Botingiz yaratgan test", 0)
    for i, q in enumerate(savollar, 1):
        doc.add_paragraph(f"{i}. {q['savol']}", style='List Number')
        for j, opt in enumerate(q['variantlar']):
            doc.add_paragraph(f"   {['A)', 'B)', 'C)', 'D)'][j]} {opt}")
            
    doc.add_heading("Javoblar kaliti:", level=2)
    kalit_matni = " | ".join([f"{i}-" + ["A", "B", "C", "D"][q['togri_index']] for i, q in enumerate(savollar, 1)])
    doc.add_paragraph(kalit_matni)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    document = BufferedInputFile(file_stream.read(), filename=f"TestTuzar_{quiz_id}.docx")
    await bot.send_document(chat_id=call.message.chat.id, document=document, caption="📄 Mana, toza test formati!")
    await call.answer()


# --- ADMIN BUYRUQLARI ---
@dp.message(Command("admin"))
async def admin_panel(message: types.Message):
    if message.from_user.id != int(ADMIN_ID): return
    await message.answer("👑 **Admin Panelga Xush Kelibsiz!**", reply_markup=admin_menyu, parse_mode="Markdown")

@dp.message(Command("profil"))
async def admin_get_profile(message: types.Message, command: CommandObject):
    if message.from_user.id != int(ADMIN_ID): return
    if not command.args:
        return await message.answer("⚠️ Foydalanuvchi ID sini kiriting. Namuna: `/profil 123456789`", parse_mode="Markdown")
        
    user_id = command.args.strip()
    async with db_pool.acquire() as conn:
        u = await conn.fetchrow("SELECT * FROM users WHERE user_id = $1", user_id)
        
    if not u:
        return await message.answer("❌ Bunday foydalanuvchi bazada topilmadi.")
        
    phone = u.get('phone_number') or "Kiritilmagan"
    text = (f"👤 **Foydalanuvchi Ma'lumotlari:**\n\n"
            f"🔹 **Ism:** {u['name']}\n"
            f"🔹 **ID:** `{u['user_id']}`\n"
            f"📱 **Telefon:** {phone}\n"
            f"🎯 **Yig'ilgan Ball:** {u['score']}\n"
            f"✅ **Yechgan testlari:** {u['tests_taken']}\n\n"
            f"🔗 **To'g'ridan-to'g'ri bog'lanish uchun pastdagi havolani bosing:**\n"
            f"👉 [Profilga o'tish](tg://user?id={u['user_id']})")
    await message.answer(text, parse_mode="Markdown")

@dp.message(F.text == "📊 Umumiy Statistika")
async def show_stats_admin(message: types.Message):
    if message.from_user.id != int(ADMIN_ID): return
    async with db_pool.acquire() as conn:
        users_count = await conn.fetchval("SELECT COUNT(*) FROM users") or 0
        quizzes_count = await conn.fetchval("SELECT COUNT(*) FROM quizzes") or 0
        image_tests = await conn.fetchval("SELECT SUM(image_tests_made) FROM users") or 0
        file_tests = await conn.fetchval("SELECT SUM(file_tests_made) FROM users") or 0
        topic_tests = await conn.fetchval("SELECT SUM(topic_tests_made) FROM users") or 0
    text = f"📊 **STARTAP STATISTIKASI:**\n\n👥 Jami a'zolar: {users_count} ta\n📝 Yaratilgan testlar: {quizzes_count} ta\n\n📈 **Tahlil:**\n📸 Rasmdan: {image_tests} marta\n📄 Fayldan: {file_tests} marta\n🧠 Mavzudan: {topic_tests} marta"
    await message.answer(text, parse_mode="Markdown")

@dp.message(F.text == "👥 Foydalanuvchilar")
async def get_users_list(message: types.Message):
    if message.from_user.id != int(ADMIN_ID): return
    wait_msg = await message.answer("⏳ Ro'yxat yuklanmoqda...")
    async with db_pool.acquire() as conn:
        users = await conn.fetch("SELECT user_id, name, phone_number, COALESCE(score, 0) as score, tests_taken FROM users ORDER BY joined_at ASC")
    if not users: return await wait_msg.edit_text("Bazada hali foydalanuvchilar yo'q.")
    
    text_content = "TESTTUZAR - FOYDALANUVCHILAR RO'YXATI (Eng eskilari birinchi)\n============================================================\n\n"
    for i, u in enumerate(users, 1):
        tel = f" | Tel: {u['phone_number']}" if u.get('phone_number') else ""
        text_content += f"{i}. Ism: {u['name'] or 'Noma`lum'} | ID: {u['user_id']}{tel} | Ball: {u['score']} | Testlar: {u['tests_taken']}\n"
    file_stream = io.BytesIO(text_content.encode('utf-8'))
    document = BufferedInputFile(file_stream.read(), filename="foydalanuvchilar_royxati.txt")
    await wait_msg.delete()
    await bot.send_document(chat_id=message.chat.id, document=document, caption=f"👥 Jami: {len(users)} ta a'zo ro'yxati.")

@dp.message(F.text == "📣 Xabar tarqatish")
async def ask_broadcast(message: types.Message, state: FSMContext):
    if message.from_user.id != int(ADMIN_ID): return
    await message.answer("Barchaga xabarni kiriting:", reply_markup=ReplyKeyboardMarkup(keyboard=[[KeyboardButton(text="🔙 Bosh menyu")]], resize_keyboard=True))
    await state.set_state(AdminState.xabar_kutish)

@dp.message(AdminState.xabar_kutish)
async def send_broadcast(message: types.Message, state: FSMContext):
    if message.text == "🔙 Bosh menyu":
        await state.clear()
        return await message.answer("Bekor qilindi.", reply_markup=asosiy_menyu)
    async with db_pool.acquire() as conn: users = await conn.fetch("SELECT user_id FROM users")
    await message.answer("⏳ Xabar tarqatilmoqda...")
    success = 0
    for u in users:
        try:
            await bot.copy_message(chat_id=int(u['user_id']), from_chat_id=message.chat.id, message_id=message.message_id)
            success += 1
            await asyncio.sleep(0.05)
        except Exception: pass
    await message.answer(f"✅ Yetib bordi: {success} ta", reply_markup=asosiy_menyu)
    await state.clear()


# --- ASOSIY BUYRUQLAR VA TEST YECHISH ---
@dp.message(F.text == "🔙 Bosh menyu")
async def back_to_main_from_admin(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer("Asosiy menyuga qaytdingiz.", reply_markup=asosiy_menyu)

@dp.message(F.text == "🔙 Bekor qilish")
async def bekor_qilish(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer("❌ Jarayon bekor qilindi.", reply_markup=asosiy_menyu)

@dp.message(Command("stop"))
async def stop_quiz_command(message: types.Message):
    stopped = False
    chat_id = message.chat.id
    user_id = message.from_user.id
    
    if ACTIVE_TESTS.get(chat_id):
        ACTIVE_TESTS[chat_id] = False
        stopped = True
        
    if ACTIVE_TESTS.get(user_id):
        ACTIVE_TESTS[user_id] = False
        if user_id in USER_EVENTS:
            USER_EVENTS[user_id].set() 
        stopped = True
        
    if stopped:
        await message.answer("🛑 Test muvaffaqiyatli to'xtatildi!", reply_markup=asosiy_menyu)
    else:
        await message.answer("⚠️ Hozirda faol test mavjud emas.", reply_markup=asosiy_menyu)

@dp.message(Command("start"))
async def start(message: types.Message, state: FSMContext, command: CommandObject = None):
    await state.clear()
    await add_user(message.from_user.id, message.from_user.first_name)

    if command and command.args:
        quiz_id = command.args
        async with db_pool.acquire() as conn:
            quiz_row = await conn.fetchrow("SELECT savollar, timer FROM quizzes WHERE quiz_id = $1", quiz_id)
            
            if quiz_row:
                await conn.execute("UPDATE users SET tests_taken = tests_taken + 1 WHERE user_id = $1", str(message.from_user.id))
                taymer = quiz_row.get('timer') or 45 
                
                await message.answer(f"🚀 Test boshlanmoqda... (Taymer: {taymer} soniya)\nTo'xtatish uchun istalgan vaqtda /stop ni bosing.", reply_markup=ReplyKeyboardRemove())
                savollar = json.loads(quiz_row['savollar'])
                
                # YANGILIK: Sessiya ballarini nolga tushirish
                SESSION_SCORES[message.from_user.id] = 0 
                USER_EVENTS[message.from_user.id] = asyncio.Event()
                ACTIVE_TESTS[message.from_user.id] = True 
                
                for i, data in enumerate(savollar, 1):
                    if not ACTIVE_TESTS.get(message.from_user.id, True):
                        break 

                    q = f"[{i}/{len(savollar)}] {data['savol'][:200]}"
                    opts = [str(opt)[:100] for opt in data['variantlar']][:4]
                    correct = int(data.get('togri_index', 0))
                    
                    sent_poll = await bot.send_poll(
                        chat_id=message.chat.id, question=q, options=opts, type='quiz', 
                        correct_option_id=correct, is_anonymous=False, open_period=taymer
                    )
                    POLL_DATA[sent_poll.poll.id] = {"correct": correct, "points": 2}
                    
                    USER_EVENTS[message.from_user.id].clear()
                    try: await asyncio.wait_for(USER_EVENTS[message.from_user.id].wait(), timeout=taymer)
                    except asyncio.TimeoutError: pass
                    await asyncio.sleep(0.5)
                    
                ACTIVE_TESTS[message.from_user.id] = False
                
                # YANGILIK: Test oxirida aniq nechta topganini e'lon qilish
                correct_ans = SESSION_SCORES.get(message.from_user.id, 0)
                await message.answer(f"🏁 **Test yakuniga yetdi!**\n\n📊 Natijangiz: {len(savollar)} ta savoldan **{correct_ans} tasiga** to'g'ri javob berdingiz! 🎯", reply_markup=asosiy_menyu, parse_mode="Markdown")
                return
            else:
                return await message.answer("⚠️ Bu test eskirgan yoki topilmadi.", reply_markup=asosiy_menyu)
                
    async with db_pool.acquire() as conn:
        user_record = await conn.fetchrow("SELECT phone_number FROM users WHERE user_id = $1", str(message.from_user.id))
    
    if user_record and user_record['phone_number']:
        await message.answer("Assalomu alaykum! EdTech platformamizga xush kelibsiz.\n\n📸 Shunchaki daftaringizni rasmga oling, PDF fayl yuboring yoki mavzu yozing!", reply_markup=asosiy_menyu)
    else:
        await message.answer("Assalomu alaykum! EdTech platformamizga xush kelibsiz.\n\n⚙️ Tizimdan qulayroq foydalanish uchun telefon raqamingizni ulashishingiz mumkin (Majburiy emas):", reply_markup=ixtiyoriy_raqam_menyu)

@dp.poll_answer()
async def handle_poll_answer(poll_answer: types.PollAnswer):
    poll_id = poll_answer.poll_id
    user_id_int = poll_answer.user.id
    tanlangan_javob = poll_answer.option_ids[0] if poll_answer.option_ids else -1
    
    if poll_id in POLL_DATA and tanlangan_javob == POLL_DATA[poll_id]["correct"]:
        ball = POLL_DATA[poll_id]["points"]
        async with db_pool.acquire() as conn:
            await conn.execute("UPDATE users SET score = COALESCE(score, 0) + $1 WHERE user_id = $2", ball, str(user_id_int))
            
        # YANGILIK: To'g'ri javoblar sonini sessiyada saqlab borish
        if user_id_int in SESSION_SCORES:
            SESSION_SCORES[user_id_int] += 1
            
    if user_id_int in USER_EVENTS:
        USER_EVENTS[user_id_int].set()

@dp.message(F.text == "📊 Mening natijalarim")
async def show_profile(message: types.Message):
    async with db_pool.acquire() as conn:
        u = await conn.fetchrow("SELECT COALESCE(score, 0) as score, tests_taken, image_tests_made, file_tests_made FROM users WHERE user_id = $1", str(message.from_user.id))
    if u: 
        text = f"📊 **Shaxsiy Statistika:**\n🎯 Yig'ilgan ball: {u['score']}\n✅ Yechilgan testlar: {u['tests_taken']}\n\n🛠 **Siz yaratgan testlar:**\n📸 Rasmdan: {u['image_tests_made']} ta\n📄 Fayldan: {u['file_tests_made']} ta"
        await message.answer(text, parse_mode="Markdown")
    else: 
        await message.answer("Siz hali bazada yo'qsiz. /start ni bosing.")

# YANGILIK: Aqlli Reyting tizimi
@dp.message(F.text == "🏆 Reyting")
async def show_reyting(message: types.Message):
    async with db_pool.acquire() as conn:
        top_users = await conn.fetch("SELECT user_id, name, COALESCE(score, 0) as score FROM users ORDER BY score DESC LIMIT 10")
        me = await conn.fetchrow("SELECT COALESCE(score, 0) as score FROM users WHERE user_id = $1", str(message.from_user.id))

    text = "🏆 **TOP-10 QAHRAMONLAR:**\n\n"
    first_place_score = top_users[0]['score'] if top_users else 0
    
    for i, u in enumerate(top_users, 1): 
        ism = u['name'] if u['name'] else "A'zo"
        # O'zini ro'yxatda belgilab ko'rsatish
        if str(u['user_id']) == str(message.from_user.id):
            ism = "👉 " + ism 
        text += f"{i}. {ism} — {u['score']} ball\n"
        
    if me:
        my_score = me['score']
        if my_score < first_place_score:
            diff = first_place_score - my_score
            # 1 ta to'g'ri javob 2 ball beradi, shuning uchun 2 ga bo'lamiz
            needed_answers = (diff // 2) + (diff % 2) 
            text += f"\n💡 **Ma'lumot:** 1-o'ringa chiqish uchun sizga yana kamida **{needed_answers} ta to'g'ri javob** kerak! Olg'a! 🚀"
        elif my_score == first_place_score and first_place_score > 0:
            text += f"\n💡 **Ma'lumot:** Tabriklaymiz, hozirda peshqadamsiz! 🥇"
            
    await message.answer(text, parse_mode="Markdown")

# --- 1-CLICK TEST TIZIMI ---
@dp.message(F.text == "📸 Rasmdan test")
async def ask_photo(message: types.Message): await message.answer("📸 Lug'at daftaringizni aniq rasmga olib yuboring.")

@dp.message(F.text == "📚 Matn/Mavzudan test")
async def ask_topic(message: types.Message): await message.answer("📄 PDF/Word fayl tashlang YOKI biror mavzuni yozing.")

@dp.message(F.photo)
async def auto_photo_handler(message: types.Message, state: FSMContext):
    await state.update_data(source_type='image', payload=message.photo[-1].file_id)
    await state.set_state(QuickQuizForm.soni)
    await message.answer("📸 Rasm qabul qilindi! Nechta savol tuzamiz?", reply_markup=soni_menyu)

@dp.message(F.document)
async def auto_doc_handler(message: types.Message, state: FSMContext):
    if not (message.document.file_name.endswith('.pdf') or message.document.file_name.endswith('.docx')): return await message.answer("⚠️ Faqat PDF yoki Word fayllar.")
    await state.update_data(source_type='file', payload=message.document.file_id, filename=message.document.file_name)
    await state.set_state(QuickQuizForm.soni)
    await message.answer("📄 Fayl qabul qilindi! Nechta savol tuzamiz?", reply_markup=soni_menyu)

@dp.message(StateFilter(None), F.text, ~F.text.in_(["📸 Rasmdan test", "📚 Matn/Mavzudan test", "📊 Mening natijalarim", "🏆 Reyting", "🔙 Bekor qilish", "/start", "/stop", "/admin", "💬 Taklif va Xatolar", "➡️ Asosiy menyuga o'tish", "🇺🇿 O'zbek tili", "🇷🇺 Русский", "🇬🇧 English"]))
async def auto_topic_handler(message: types.Message, state: FSMContext):
    if message.text.isdigit(): return
    await state.update_data(source_type='topic', payload=message.text)
    await state.set_state(QuickQuizForm.soni)
    await message.answer("🧠 Mavzu qabul qilindi! Nechta savol tuzamiz?", reply_markup=soni_menyu)

@dp.message(QuickQuizForm.soni)
async def ask_lang_handler(message: types.Message, state: FSMContext):
    if not message.text or not message.text.isdigit(): 
        return await message.answer("⚠️ Iltimos, pastdagi tugmalardan sonni tanlang.", reply_markup=soni_menyu)
        
    await state.update_data(soni=int(message.text))
    await state.set_state(QuickQuizForm.til) 
    await message.answer("🌐 Qaysi tilda test tuzamiz?", reply_markup=til_menyu)

@dp.message(QuickQuizForm.til)
async def ask_timer_handler(message: types.Message, state: FSMContext):
    tanlangan_til = message.text
    if tanlangan_til not in ["🇺🇿 O'zbek tili", "🇷🇺 Русский", "🇬🇧 English"]:
        return await message.answer("⚠️ Iltimos, pastdagi tugmalardan tilni tanlang.", reply_markup=til_menyu)
        
    if tanlangan_til == "🇺🇿 O'zbek tili":
        til_nomi = "Uzbek"
    elif tanlangan_til == "🇷🇺 Русский":
        til_nomi = "Russian"
    else:
        til_nomi = "English"
        
    await state.update_data(til=til_nomi)
    await state.set_state(QuickQuizForm.vaqt)
    await message.answer("⏱ Har bir savol uchun qancha vaqt ajratamiz?", reply_markup=vaqt_menyu)

@dp.message(QuickQuizForm.vaqt)
async def generate_magic(message: types.Message, state: FSMContext):
    vaqt_matni = message.text.replace("soniya", "").strip()
    if not vaqt_matni.isdigit(): return await message.answer("⚠️ Iltimos, tugmalardan vaqtni tanlang.", reply_markup=vaqt_menyu)
        
    tanlangan_vaqt = int(vaqt_matni)
    data = await state.get_data()
    soni, source, til_nomi = data['soni'], data['source_type'], data['til']
    
    wait_msg = await message.answer("⚙️ Sun'iy intellekt tahlil qilmoqda... Kuting.", reply_markup=ReplyKeyboardRemove())
    
    qatiy_buyruq = f"DIQQAT: Vazifang faqat va faqat berilgan matn/mavzu doirasida {soni} ta test tuzish. Mavzudan umuman tashqariga chiqma. Test tili: {til_nomi}. Agar matnda yetarli ma'lumot bo'lmasa, o'zingdan to'qima. FAQAT JSON ARRAY qaytar. Namuna: [{{\"savol\": \"...\", \"variantlar\": [\"A\", \"B\", \"C\", \"D\"], \"togri_index\": 0}}]"

    try:
        if source == 'image':
            file_info = await bot.get_file(data['payload'])
            file_data = io.BytesIO()
            await bot.download_file(file_info.file_path, destination=file_data)
            file_data.seek(0)
            img = Image.open(file_data).convert("RGB")
            prompt = f"{qatiy_buyruq}\nRasmdagi matnlarni diqqat bilan o'qib, shunga doir test tuz."
            response = await asyncio.to_thread(model.generate_content, [prompt, img])
            
        elif source == 'file':
            file_info = await bot.get_file(data['payload'])
            file_data = io.BytesIO()
            await bot.download_file(file_info.file_path, destination=file_data)
            file_data.seek(0)
            text = await asyncio.to_thread(read_file_sync, file_data, data.get('filename', ''))
            prompt = f"{qatiy_buyruq}\n\nMatn: {text[:15000]}"
            response = await asyncio.to_thread(model.generate_content, prompt)
            
        elif source == 'topic':
            prompt = f"{qatiy_buyruq}\n\nMavzu: '{data['payload']}'"
            response = await asyncio.to_thread(model.generate_content, prompt)

        json_matn = clean_json_text(response.text)
        savollar = json.loads(json_matn)
        
        for q in savollar:
            eski_index = int(q.get('togri_index', 0))
            if eski_index < 0 or eski_index > 3: eski_index = 0
            togri_matn = q['variantlar'][eski_index]
            random.shuffle(q['variantlar'])
            q['togri_index'] = q['variantlar'].index(togri_matn)

        quiz_id = str(uuid.uuid4())[:8]
        async with db_pool.acquire() as conn:
            await conn.execute("INSERT INTO quizzes (quiz_id, source_type, savollar, timer) VALUES ($1, $2, $3, $4)", quiz_id, source, json.dumps(savollar), tanlangan_vaqt)
            if source == 'image': await conn.execute("UPDATE users SET image_tests_made = image_tests_made + 1 WHERE user_id = $1", str(message.from_user.id))
            elif source == 'file': await conn.execute("UPDATE users SET file_tests_made = file_tests_made + 1 WHERE user_id = $1", str(message.from_user.id))
            elif source == 'topic': await conn.execute("UPDATE users SET topic_tests_made = topic_tests_made + 1 WHERE user_id = $1", str(message.from_user.id))

        bot_info = await bot.get_me()
        test_link = f"https://t.me/{bot_info.username}?start={quiz_id}"
        
        inline_kb = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="🚀 Testni boshlash", url=test_link)], 
            [InlineKeyboardButton(text="📥 Fayl qilib olish (Word)", callback_data=f"down_{quiz_id}")]
        ])
        
        await wait_msg.delete()
        msg_text = f"✅ **Testingiz tayyor! ({len(savollar)} ta savol)**\n*(Har biriga {tanlangan_vaqt} soniya)*\n\n👥 **Guruhda o'ynash uchun:** Botni guruhingizga qo'shing va u yerga `/quiz {quiz_id}` deb yozing!"
        await message.answer(msg_text, reply_markup=inline_kb, parse_mode="Markdown")
        await message.answer("Asosiy menyuga qaytdingiz.", reply_markup=asosiy_menyu)
        await state.clear()

    except Exception as e:
        with suppress(Exception): await wait_msg.delete()
        print(f"XATOLIK YUZ BERDI: {e}") 
        await message.answer("⚠️ Sun'iy intellekt xato qildi yoki savol tuza olmadi. Boshqatdan urinib ko'ring.", reply_markup=asosiy_menyu)
        await state.clear()


async def main():
    keep_alive()
    try:
        await init_db_pool()
        print("✅ Baza muvaffaqiyatli ulandi!")
    except Exception as e:
        print(f"❌ BAZAGA ULANISHDA XATOLIK: {e}")
        return
        
    print("🚀 UPPERLAR MVP BOTI ISHGA TUSHDI!")
    try:
        await bot.delete_webhook(drop_pending_updates=True)
        await dp.start_polling(bot)
    except Exception as e:
        print(f"❌ BOT MOTORIDA XATO: {e}")

if __name__ == "__main__": 
    asyncio.run(main())
